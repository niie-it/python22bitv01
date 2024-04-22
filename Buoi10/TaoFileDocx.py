# TaoFileDocx.py
# Phải cài thư viện python-docx trước
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

document.add_heading("Chương 1: Mở đầu", 1)
document.add_paragraph("Giới thiệu sơ lược về đề tài: Phục vụ cho ai? Làm gì?")
document.add_heading("1.2: Công nghệ sử dụng", 2)

records = [
    {"mon":"Lập trình nâng cao", "so_tc": 3, "gpa": 8.9},
    {"mon":"Bảo mật", "so_tc": 3, "gpa": 9.2},
    {"mon":"Tiếng Anh thương mại", "so_tc": 4, "gpa": 7.9}
]

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Môn'
hdr_cells[1].text = 'Số tín chỉ'
hdr_cells[2].text = 'Điểm'
for item in records:
    row_cells = table.add_row().cells
    row_cells[0].text = item["mon"]
    row_cells[1].text = str(item["so_tc"])
    row_cells[2].text = str(item["gpa"])

document.save("MyDocx.docx")